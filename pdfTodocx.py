from pdf2docx import Converter
from PyPDF2 import PdfFileReader, PdfFileWriter
from PyPDF2.pdf import ContentStream
from PyPDF2.generic import TextStringObject, NameObject
from PyPDF2.utils import b_
import docx
from docx import Document 


pdf_file = r'C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\Tesik_Loader.pdf'
docx_file = r'C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\TesikLoader.docx'

# convert pdf to docx
cv = Converter(pdf_file)
cv.convert(docx_file, start=0, end=None)
cv.close()

# path = r'C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\sample.docx'
# path1 = r'C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\sample1.docx'
# myList = []
# def getText(filename):
#     doc = docx.Document(filename)
#     fullText = []
#     for para in doc.paragraphs:
#         fullText.append(para.text)
#     #return '\n'.join(fullText)
#     return fullText
# myList = getText(path)

# def addText(file2):
#     mydoc = docx.Document()
#     for text in file2:
#         mydoc.add_paragraph(text)
#         mydoc.save(path1)
# addText(myList)
# print('hello world')