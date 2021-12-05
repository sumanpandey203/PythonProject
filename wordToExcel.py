# Import the required library docx 
import matplotlib
import docx
from textblob import TextBlob
from docx import Document 
from openpyxl import Workbook
from docx.enum.style import WD_STYLE_TYPE

##############################
myOddList = []
myEvenList = []
myTextList = []
myEven_even = []
myEven_odd= []
myEven_odd_en = []

#############################

#open excel
wb = Workbook()
sheet = wb.active
sheet.title = 'Tesiko_CR'
sheet1 = wb.create_sheet(title='Tesiko_Info')
sheet2 = wb.create_sheet(title='Tesiko_Katalyst')

#############################
 
# Specify the path of the file 
path = r'C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\TesikLoader_Work.docx'
path1 = (r"C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\TesikLoader_en1.docx")
myList = []
# read file 
document = Document(path) 
myDoc = Document()
styles = document.styles
# to know the formate of the docx
paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
for style in paragraph_styles:
    print(style.name)
print(document.styles)
counter = 0
for para in document.paragraphs:
    # print(f'{para.style.name}= {para.text}')
    if para.style.name == 'Heading 1':
        counter += 1
        myOddList.append(para.text)
    else:
        myTextList.append(para.text)
#removing space from the docx       
def removeSpace(space):
    spaceList =[]
    for val in space:
        if val != '':
            spaceList.append(val)
    return spaceList

myEvenList = removeSpace(myTextList)


print(f'count:{counter}= lenOfOddList:{len(myOddList)}')
# Translation fron deutsch to english
def transToEng(var):
    mytest = []    
    for i in var:
        if i != '':
           
            text = TextBlob(i)
            print(f'{text.detect_language()}={i}')
            if text.detect_language() != 'en':
                result = text.translate(to='en')
                mytest.append(result)
            else:
                mytest.append(i)

        else:
            mytest.append('')
            
    return mytest

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    #return '\n'.join(fullText)
    return fullText
    
for r in range(0,len(myEvenList)):   
    if r %2 == 0:
        myEven_even.append(myEvenList[r])
    else:
         myEven_odd.append(myEvenList[r])
myEven_odd_en = transToEng(myEven_odd) 


myList =[]    
myList = getText(path1)
#To check
for i in range(len(myEven_odd_en)):
    print(f'{i} = {myEven_odd_en[i]}')
# Write to the excel
for r in range(0,len(myOddList)):
    c1 = sheet.cell(row= r+1 , column = 1) 
    c1.value = myOddList[r]

# print(myEven_even)
for r in range(0,len(myEven_even)):
    c2 = sheet.cell(row= r+1 , column = 2) 
    c2.value = myEven_even[r]
for r in range(0,len(myEven_odd)):
    c3 = sheet.cell(row= r+1 , column = 3) 
    c3.value = myEven_odd[r]

for r in range(0,len(myList)):
    c1 = sheet.cell(row= r+1 , column = 4) 
    c1.value = myList[r]


wb.save(r'C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\Tesiko_22_02_21.xlsx')






