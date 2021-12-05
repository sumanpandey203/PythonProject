from textblob import TextBlob
import docx
from docx import Document 
from openpyxl import Workbook
from openpyxl import load_workbook


#path = r'C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\TesikLoader.docx'
#path1 = r'C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\TesikLoader_en.docx'
path = r"C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\ChangeRequest\ORS VEDS Technisches Sicherheitskonzept_2021_03_02.docx"

myDoc1 = Document(path)
myDoc = Document()

myList = []
myTrans =[]

# myText = open(path, 'r')
count = 0

for text in myDoc1.paragraphs:
    print(text.text)
    myList.append(text.text)

for text in myList:
    myDoc.add_paragraph(text)
    myDoc.save(r"C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\ChangeRequest\testmyText.docx")

# wb = Workbook()
# Sheet = wb.create_sheet(title='MindMap')

# for idx in range(len(myList)):
#     cell = Sheet.cell(row = idx+1, column = 1)
#     cell.value= myList[idx]
# wb.save(r"C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\ChangeRequest\testmyText.xlsx")

# def getText(filename):
#     doc = docx.Document(filename)
#     fullText = []
#     for para in doc.paragraphs:
#         fullText.append(para.text)
#     #return '\n'.join(fullText)
#     return fullText
# myList = getText(path)
# print('\n'.join(myList))
# def transToEng(var):
#     mytest = []    
#     for i in var:
#         if i != '':
#             mytest.append(i)         
#     return mytest
# myTransDoc = transToEng(myList)

# for text in myTransDoc:
#     myDoc.add_paragraph(text)
#     myDoc.save(r"C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\TesikLoader_en.docx")