from translate import Translator
import math
#from googletrans import Translator 
from textblob import TextBlob
import textblob.exceptions
import docx
from docx import Document 
from openpyxl import Workbook
from docx.enum.style import WD_STYLE_TYPE

#-------------------------------------------------------------------------------------------------
##############################
myOddList = []
myEvenList = []
myTextList = []
myEven_even = []
myEven_odd= []
myEven_odd_en = []
myOrginalDoc = []

#############################

# blob = TextBlob('Das ist ein auto.')
# print(blob.translate(to='en'))

 
# Specify the path of the file 
path = r'C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\TesikLoader_Work.docx'
#path1 = (r"C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\TesikLoader_en1.docx")
myList = []
# read file 
document = Document(path) 
myDoc = Document()
styles = document.styles
# to know the formate of the docx
paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
for style in paragraph_styles:
    print(style.name)
print(document.styles)
counter = 0
for para in document.paragraphs:
    # print(f'{para.style.name}= {para.text}')
   
    if para.style.name == 'Heading 1':
        counter += 1
        myOddList.append(para.text)
    else:
        myTextList.append(para.text)
#removing space from the docx 
def removeSpace(space):
    spaceList =[]
    for val in space:
        if val != '':
            spaceList.append(val)
    return spaceList

myEvenList = removeSpace(myTextList)      
########################################################


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    #return '\n'.join(fullText)
    return fullText
myList = getText(path)

#print('\n'.join(myList))


def transToEng(var):
    mytest = []
    txt=" "
    
    for i in var:
        if i != '':
            try:
                print(i)
                text = TextBlob(i)
                print(text.detect_language())
                if text.detect_language() != 'en' or text.detect_language() == 'ca' :
                    result = text.translate(to='en')
                    mytest.append(result)
                else:
                    mytest.append(i)
            except textblob.exceptions.TranslatorError:
                word=" "  #replace word with space
                txt=txt+i
                mytest.append(i)
            except textblob.exceptions.NotTranslated:
                txt=txt+i+" "
                mytest.append(i)
            # else:
            #     txt=txt+i+" "
            #     mytest.append(i)

        else:
            mytest.append('')
            
    return mytest

myEvenList = removeSpace(myList)   

myOrginalDoc = transToEng(myList)
print(myOrginalDoc)
#myTrans = getText(path1)



# for index in range(len(myList)):
#     myDoc.add_paragraph(myList[index])
#     myDoc.add_paragraph(myOrginalDoc[index])
#     myDoc.save(r"C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\TesikCombi_DE_EN.docx") 


   







