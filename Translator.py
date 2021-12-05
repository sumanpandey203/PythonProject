from translate import Translator
import math
#from googletrans import Translator 
from textblob import TextBlob
import docx
from docx import Document 
from openpyxl import Workbook
from docx.enum.style import WD_STYLE_TYPE

##############################
myOddList = []
myEvenList = []
myTextList = []
myEven_even = []
myEven_odd= []
myEven_odd_en = []

#############################

blob = TextBlob('Das ist ein auto.')
print(blob.translate(to='en'))

 
# Specify the path of the file 
path = r'C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\TesikLoader_Work.docx'
#path1 = (r"C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\TesikLoader_en1.docx")
myList = []
# read file 
document = Document(path) 
myDoc = Document()
styles = document.styles
# to know the formate of the docx
paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
for style in paragraph_styles:
    print(style.name)
print(document.styles)
counter = 0
for para in document.paragraphs:
    # print(f'{para.style.name}= {para.text}')
    if para.style.name == 'Heading 1':
        counter += 1
        myOddList.append(para.text)
    else:
        myTextList.append(para.text)
#removing space from the docx       
def removeSpace(space):
    spaceList =[]
    for val in space:
        if val != '':
            spaceList.append(val)
    return spaceList

myEvenList = removeSpace(myTextList)

for r in range(0,len(myEvenList)):   
    if r %2 == 0:
        myEven_even.append(myEvenList[r])
    else:
         myEven_odd.append(myEvenList[r])

########################################################


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    #return '\n'.join(fullText)
    return fullText
myList = getText(path)

print('\n'.join(myList))
def transToEng(var):
    mytest = []
    
    for i in var:
        if i != '':
            print(i)
            text = TextBlob(i)
            print(text.detect_language())
            if text.detect_language() != 'en':
                result = text.translate(to='en')
                mytest.append(result)
            else:
                mytest.append(i)

        else:
            mytest.append('')
            
    return mytest
myTransDoc = transToEng(myEven_odd_en)
myEven_odd_en = transToEng(myEven_odd) 
#myTrans = getText(path1)



for text in myEven_odd_en:
    myDoc.add_paragraph(text)
    myDoc.save(r"C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\TesikLoader_en2.docx")

# for r in range(0,len(myList)):
#     c1 = sheet3.cell(row= r+1 , column = 1) 
#     c1.value = myList[r]
# for r in range(0,len(myTrans)):
#     c2 = sheet3.cell(row= r+1 , column = 2) 
#     c2.value = myTrans[r]

# wb.save(r'C:\Users\suman.pandey\Desktop\work_2021\changeRequest_21\demo.xlsx')

   






